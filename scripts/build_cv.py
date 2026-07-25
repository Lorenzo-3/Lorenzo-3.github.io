#!/usr/bin/env python3
"""Build the English and Italian CV PDFs' source DOCX files."""

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "assets" / "cv"

INK = RGBColor(27, 23, 16)
MUTED = RGBColor(91, 79, 61)
TEAL = RGBColor(0, 90, 86)
GOLD = RGBColor(150, 110, 0)
MAROON = RGBColor(193, 0, 48)
PALE_GOLD = "FBF6DE"
PALE_TEAL = "E8F3F1"
LIGHT_BORDER = "D9D1BC"


CONTENT = {
    "en": {
        "filename": "Lorenzo_Marinelli_CV_EN.docx",
        "cv_label": "CURRICULUM VITAE",
        "date": "JUNE 2026",
        "subtitle": "MSc Computer Science Student | Applied AI & Scientific Computing",
        "location": "Italy",
        "profile_heading": "Profile",
        "profile": (
            "Computer Science MSc student at Sapienza University of Rome working across applied AI, "
            "scientific data, and research-oriented software. I build reproducible experiments and "
            "interactive technical tools spanning graph machine learning, bioinformatics, computer "
            "vision, time series, reinforcement learning, and language-model-guided workflows."
        ),
        "education": "Education",
        "education_items": [
            ("MSc Computer Science", "2025-Present", "Sapienza University of Rome", None),
            (
                "Bachelor's Degree in Applied Computer Science & Artificial Intelligence",
                "2022-2025",
                "Sapienza University of Rome",
                "Final grade: 110 cum laude",
            ),
            (
                "Scientific High School Diploma",
                "2017-2022",
                "Liceo Scientifico A. Romita",
                "Final grade: 100 e lode",
            ),
        ],
        "technical": "Technical Profile",
        "skill_items": [
            ("Programming & data", "Python, NumPy, pandas, Matplotlib, scikit-learn, notebooks, SQL"),
            ("Deep learning", "PyTorch, PyTorch Lightning, PyTorch Geometric, torchvision, model evaluation"),
            ("Vision & graphs", "OpenCV, NetworkX, Gephi, anomaly localization, graph analysis"),
            ("Engineering", "Git, Docker, HTML/CSS, JavaScript, OpenAPI, lightweight web/API development"),
            ("Research workflow", "Reproducible runs, visual diagnostics, documented experiments, compact demos"),
        ],
        "focus": "Current Focus & Availability",
        "focus_items": [
            "Open to internships, thesis work, research prototypes, and open-source collaboration.",
            "Primary interests: deep learning, graph ML, bioinformatics, data analysis, and ML security.",
            "Working style: start from the question and data, then choose models and evaluate limitations clearly.",
        ],
        "projects": "Selected Project Experience",
        "thesis_title": "Bachelor's Thesis - GNN Protein Interaction Prediction",
        "thesis_url": "https://www.lmarinelli.eu/projects/thesis-gnn-protein-interactions/",
        "thesis_meta": "Bioinformatics | Graph ML | Link prediction | ProtT5 embeddings",
        "thesis_bullets": [
            "Built a graph-neural-network workflow to rank likely missing protein-protein interactions, combining protein sequence information with known interaction topology and a NOTCH2 case study.",
            "Used TAGConv, TransformerConv, and GINConv encoding with a pair decoder; evaluated ranking quality under a 1:10 positive-negative skew.",
            "Results: AUROC 0.96, AUPRC 0.89, Precision@500 1.00, about 81% recall at threshold 0.5; documented topology bias and scalability limits.",
        ],
        "thesis_metrics": "THESIS RESULTS   AUROC 0.96   |   AUPRC 0.89   |   P@500 1.00   |   Recall ~81%",
        "rl_title": "RL-Nav v1 - Oracle-Guided RL for Partial-Observation Navigation",
        "rl_url": "https://www.lmarinelli.eu/assets/RL-Nav/RL-Nav_v1_Lorenzo_Marinelli.pdf",
        "rl_meta": "Reinforcement learning | LLM guidance | Recurrent control | PyBullet",
        "rl_bullets": [
            "Designed the interface and constrained structured outputs for an LLM oracle guiding navigation in a stochastic, partially observed environment with a hidden goal.",
            "Reached 62% held-out success and identified robust execution as the principal bottleneck.",
        ],
        "uav_title": "UAV Anomaly Detection & Localization",
        "uav_url": "https://www.lmarinelli.eu/projects/uav-anomaly-detection/",
        "uav_meta": "Computer vision | AE/VAE | Heatmaps | Connected components",
        "uav_bullets": [
            "Developed a modular proposal pipeline for high-resolution aerial imagery: reconstruct the expected scene, score residual anomalies, threshold and clean masks, then extract boxes and crops.",
            "Focused on interpretable anomaly maps and lightweight candidate generation before downstream classification.",
        ],
        "lorenz_title": "Lorenz Attractor Forecasting",
        "lorenz_url": "https://www.lmarinelli.eu/projects/lorenz/",
        "lorenz_meta": "LSTM | Time series | Chaotic systems | Numerical simulation",
        "lorenz_bullets": [
            "Compared coordinate, derivative-aware, and residual-style recurrent forecasting approaches on a chaotic dynamical system.",
            "Used recursive rollouts to study drift, stability, and error accumulation beyond short-horizon metrics.",
        ],
        "additional_projects": "Additional Projects",
        "additional_project_items": [
            (
                "EU Elections 2019 Data Analysis",
                "Cleaned public data, checked hypotheses, and produced reproducible visual and narrative analysis.",
            ),
            (
                "Synthetic Dataset Generation in Blender",
                "Generated labeled, domain-randomized renders for YOLO-style object-detection training.",
            ),
            (
                "Small Web / API Projects",
                "Built lightweight apps and API integrations using OpenAPI, SQL, HTML, CSS, and JavaScript; includes WasaText, Mood Tracker, and this portfolio.",
            ),
        ],
        "portfolio": "Interactive Technical Portfolio",
        "portfolio_intro": (
            "Built fourteen browser-native explainers, simulations, and utilities that turn technical concepts "
            "into inspectable interactions."
        ),
        "portfolio_groups": [
            (
                "Graph & machine learning",
                "Graph Signal Diffusion; GNN Message Passing Toy; Clustering Lab; MLP Decision Boundary Lab; Anomaly Node Game",
            ),
            (
                "Signals, vision & generative models",
                "Fourier Transform Playground; Signal Filter Playground; Convolution Kernel Playground; Computer Vision Mini Lab; Diffusion Model Denoising Toy",
            ),
            (
                "Scientific & creative computing",
                "Gravity Assist Sandbox; Circle of Fifths Network; Image Color Editor",
            ),
            ("Browser utility", "Local File Converter using browser-native tools and FFmpeg.wasm"),
        ],
        "playground_url": "https://www.lmarinelli.eu/playground/",
        "practice": "Research Practice",
        "practice_items": [
            "Connect coursework, thesis work, and personal projects through readable notes, demos, and practical evaluation.",
            "Turn domain problems and messy datasets into reproducible analyses, visual summaries, and documented experiments.",
            "Maintain a learning log; current reference work includes an NVIDIA DLI course on adversarial machine learning and model security.",
        ],
        "interests": "Broader Interests",
        "interests_text": (
            "Statistics, psychology, physics, biology, and philosophy; writing music, classical guitar, "
            "3D modeling, drawing, photography, films, television series, and video games."
        ),
        "links": "Portfolio & Contact",
        "links_intro": "Project details, thesis artifacts, presentations, and interactive demos are available online.",
        "link_items": [
            ("Portfolio", "https://www.lmarinelli.eu"),
            ("GitHub", "https://github.com/Lorenzo-3"),
            ("LinkedIn", "https://www.linkedin.com/in/lorenzo-marinelli-6821403a9/"),
            ("Bachelor's thesis PDF", "https://www.lmarinelli.eu/assets/thesis/thesis.pdf"),
            ("Thesis slides", "https://www.lmarinelli.eu/assets/thesis/thesis-slides.pdf"),
            ("RL-Nav presentation", "https://www.lmarinelli.eu/assets/RL-Nav/RL-Nav_v1_Lorenzo_Marinelli.pdf"),
        ],
    },
    "it": {
        "filename": "Lorenzo_Marinelli_CV_IT.docx",
        "cv_label": "CURRICULUM VITAE",
        "date": "GIUGNO 2026",
        "subtitle": "Studente Magistrale in Computer Science | AI Applicata & Calcolo Scientifico",
        "location": "Italia",
        "profile_heading": "Profilo",
        "profile": (
            "Studente magistrale in Computer Science alla Sapienza Università di Roma, attivo tra AI applicata, "
            "dati scientifici e software orientato alla ricerca. Sviluppo esperimenti riproducibili e strumenti "
            "tecnici interattivi su graph machine learning, bioinformatica, visione artificiale, serie temporali, "
            "reinforcement learning e workflow guidati da modelli linguistici."
        ),
        "education": "Formazione",
        "education_items": [
            ("Laurea Magistrale in Computer Science", "2025-Presente", "Sapienza Università di Roma", None),
            (
                "Laurea Triennale in Applied Computer Science & Artificial Intelligence",
                "2022-2025",
                "Sapienza Università di Roma",
                "Voto finale: 110 e lode",
            ),
            (
                "Diploma di Liceo Scientifico",
                "2017-2022",
                "Liceo Scientifico A. Romita",
                "Voto finale: 100 e lode",
            ),
        ],
        "technical": "Profilo Tecnico",
        "skill_items": [
            ("Programmazione & dati", "Python, NumPy, pandas, Matplotlib, scikit-learn, notebook, SQL"),
            ("Deep learning", "PyTorch, PyTorch Lightning, PyTorch Geometric, torchvision, valutazione modelli"),
            ("Visione & grafi", "OpenCV, NetworkX, Gephi, localizzazione anomalie, analisi di grafi"),
            ("Ingegneria", "Git, Docker, HTML/CSS, JavaScript, OpenAPI, sviluppo leggero web/API"),
            ("Metodo di ricerca", "Run riproducibili, diagnostica visuale, esperimenti documentati, demo compatte"),
        ],
        "focus": "Direzione Attuale & Disponibilità",
        "focus_items": [
            "Aperto a tirocini, tesi, prototipi di ricerca e collaborazioni open-source.",
            "Interessi principali: deep learning, graph ML, bioinformatica, analisi dati e sicurezza dei modelli ML.",
            "Metodo: partire dalla domanda e dai dati, poi scegliere i modelli e valutarne chiaramente i limiti.",
        ],
        "projects": "Esperienza Progettuale Selezionata",
        "thesis_title": "Tesi Triennale - Predizione di Interazioni Proteiche con GNN",
        "thesis_url": "https://www.lmarinelli.eu/it/projects/thesis-gnn-protein-interactions/",
        "thesis_meta": "Bioinformatica | Graph ML | Link prediction | Embedding ProtT5",
        "thesis_bullets": [
            "Sviluppato un workflow con graph neural network per ordinare interazioni proteina-proteina probabilmente mancanti, combinando informazione di sequenza e topologia nota con un case study su NOTCH2.",
            "Usati encoder TAGConv, TransformerConv e GINConv con pair decoder; valutata la qualità del ranking con sbilanciamento positivi-negativi 1:10.",
            "Risultati: AUROC 0.96, AUPRC 0.89, Precision@500 1.00, recall circa 81% alla soglia 0.5; documentati bias topologico e limiti di scalabilità.",
        ],
        "thesis_metrics": "RISULTATI TESI   AUROC 0.96   |   AUPRC 0.89   |   P@500 1.00   |   Recall ~81%",
        "rl_title": "RL-Nav v1 - RL Guidato da Oracolo in Osservabilità Parziale",
        "rl_url": "https://www.lmarinelli.eu/assets/RL-Nav/RL-Nav_v1_Lorenzo_Marinelli.pdf",
        "rl_meta": "Reinforcement learning | Guida LLM | Controllo ricorrente | PyBullet",
        "rl_bullets": [
            "Progettata l'interfaccia e gli output strutturati vincolati di un oracolo LLM per la navigazione in un ambiente stocastico e parzialmente osservabile con obiettivo nascosto.",
            "Raggiunto il 62% di successo in held-out, identificando la robustezza esecutiva come collo di bottiglia principale.",
        ],
        "uav_title": "Anomaly Detection & Localizzazione UAV",
        "uav_url": "https://www.lmarinelli.eu/it/projects/uav-anomaly-detection/",
        "uav_meta": "Visione artificiale | AE/VAE | Heatmap | Componenti connesse",
        "uav_bullets": [
            "Sviluppata una pipeline modulare per immagini aeree ad alta risoluzione: ricostruzione della scena attesa, scoring delle anomalie residue, pulizia delle maschere ed estrazione di box e crop.",
            "Lavoro orientato a mappe di anomalia interpretabili e generazione leggera di regioni candidate prima della classificazione.",
        ],
        "lorenz_title": "Previsione dell'Attrattore di Lorenz",
        "lorenz_url": "https://www.lmarinelli.eu/it/projects/lorenz/",
        "lorenz_meta": "LSTM | Serie temporali | Sistemi caotici | Simulazione numerica",
        "lorenz_bullets": [
            "Confrontati approcci ricorrenti di predizione delle coordinate, derivative-aware e residual-style su un sistema dinamico caotico.",
            "Usati rollout ricorsivi per studiare deriva, stabilità e accumulo degli errori oltre le metriche a breve orizzonte.",
        ],
        "additional_projects": "Altri Progetti",
        "additional_project_items": [
            (
                "Analisi Dati - Elezioni Europee 2019",
                "Pulizia di dati pubblici, verifica di ipotesi e produzione di analisi visuali e narrative riproducibili.",
            ),
            (
                "Generazione di Dataset Sintetici in Blender",
                "Generazione di render etichettati con domain randomization per training di object detection in stile YOLO.",
            ),
            (
                "Piccoli Progetti Web / API",
                "App leggere e integrazioni API con OpenAPI, SQL, HTML, CSS e JavaScript; includono WasaText, Mood Tracker e questo portfolio.",
            ),
        ],
        "portfolio": "Portfolio Tecnico Interattivo",
        "portfolio_intro": (
            "Realizzati quattordici strumenti, simulatori e spiegazioni browser-native che trasformano concetti "
            "tecnici in interazioni direttamente ispezionabili."
        ),
        "portfolio_groups": [
            (
                "Grafi & machine learning",
                "Graph Signal Diffusion; GNN Message Passing Toy; Clustering Lab; MLP Decision Boundary Lab; Anomaly Node Game",
            ),
            (
                "Segnali, visione & modelli generativi",
                "Fourier Transform Playground; Signal Filter Playground; Convolution Kernel Playground; Computer Vision Mini Lab; Diffusion Model Denoising Toy",
            ),
            (
                "Calcolo scientifico & creativo",
                "Gravity Assist Sandbox; Circle of Fifths Network; Image Color Editor",
            ),
            ("Utility browser", "File Converter locale basato su strumenti browser-native e FFmpeg.wasm"),
        ],
        "playground_url": "https://www.lmarinelli.eu/it/playground/",
        "practice": "Metodo di Ricerca",
        "practice_items": [
            "Collego corsi, tesi e progetti personali attraverso appunti leggibili, demo e valutazione pratica.",
            "Trasformo problemi di dominio e dataset disordinati in analisi riproducibili, sintesi visuali ed esperimenti documentati.",
            "Mantengo un learning log; tra i riferimenti attuali è presente un corso NVIDIA DLI su adversarial machine learning e model security.",
        ],
        "interests": "Interessi Trasversali",
        "interests_text": (
            "Statistica, psicologia, fisica, biologia e filosofia; composizione musicale, chitarra classica, "
            "modellazione 3D, disegno, fotografia, film, serie TV e videogiochi."
        ),
        "links": "Portfolio & Contatti",
        "links_intro": "Dettagli dei progetti, tesi, presentazioni e demo interattive sono disponibili online.",
        "link_items": [
            ("Portfolio", "https://www.lmarinelli.eu/it/"),
            ("GitHub", "https://github.com/Lorenzo-3"),
            ("LinkedIn", "https://www.linkedin.com/in/lorenzo-marinelli-6821403a9/"),
            ("PDF della tesi triennale", "https://www.lmarinelli.eu/assets/thesis/thesis.pdf"),
            ("Slide della tesi", "https://www.lmarinelli.eu/assets/thesis/thesis-slides.pdf"),
            ("Presentazione RL-Nav", "https://www.lmarinelli.eu/assets/RL-Nav/RL-Nav_v1_Lorenzo_Marinelli.pdf"),
        ],
    },
}


def set_run_font(run, name="Arial", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_left_border(paragraph, color, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_hyperlink(paragraph, text, url, color=TEAL, bold=False, size=9.2):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Arial")
    run_fonts.set(qn("w:hAnsi"), "Arial")
    run_properties.append(run_fonts)
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), str(color))
    run_properties.append(run_color)
    run_size = OxmlElement("w:sz")
    run_size.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(run_size)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, size=8.5, color=MUTED)


def paragraph_base(paragraph, before=0, after=4, line=1.12, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_together = True


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Georgia" if style_name in ("Title", "Heading 1") else "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)

    title = styles["Title"]
    title._element.get_or_add_pPr().remove_all("w:pBdr")
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = INK
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)

    subtitle = styles["Subtitle"]
    subtitle._element.get_or_add_pPr().remove_all("w:pBdr")
    subtitle.font.size = Pt(12.5)
    subtitle.font.bold = False
    subtitle.font.italic = False
    subtitle.font.color.rgb = TEAL
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)

    h1 = styles["Heading 1"]
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = GOLD
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(5)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.size = Pt(10.5)
    h2.font.bold = True
    h2.font.color.rgb = TEAL
    h2.paragraph_format.space_before = Pt(5)
    h2.paragraph_format.space_after = Pt(1)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.size = Pt(9.5)
    h3.font.bold = True
    h3.font.color.rgb = MAROON
    h3.paragraph_format.space_before = Pt(3)
    h3.paragraph_format.space_after = Pt(1)
    h3.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    bullet.font.size = Pt(9.3)
    bullet.font.color.rgb = INK
    bullet.paragraph_format.left_indent = Inches(0.3)
    bullet.paragraph_format.first_line_indent = Inches(-0.16)
    bullet.paragraph_format.space_after = Pt(2.5)
    bullet.paragraph_format.line_spacing = 1.1

    if "CV Meta" not in styles:
        meta = styles.add_style("CV Meta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["CV Meta"]
    meta.font.name = "Arial"
    meta._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    meta.font.size = Pt(8.7)
    meta.font.italic = True
    meta.font.color.rgb = MUTED
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.line_spacing = 1.0

    if "CV Entry" not in styles:
        entry = styles.add_style("CV Entry", WD_STYLE_TYPE.PARAGRAPH)
    else:
        entry = styles["CV Entry"]
    entry.font.name = "Arial"
    entry._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    entry._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    entry.font.size = Pt(9.5)
    entry.font.color.rgb = INK
    entry.paragraph_format.space_after = Pt(1)
    entry.paragraph_format.keep_with_next = True


def configure_page(document, data):
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph_base(header_p, after=0, line=1.0)
    run = header_p.add_run(f"LORENZO MARINELLI  |  {data['cv_label']}")
    set_run_font(run, size=7.8, color=TEAL, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    paragraph_base(footer_p, after=0, line=1.0)
    add_hyperlink(footer_p, "www.lmarinelli.eu", "https://www.lmarinelli.eu", color=MUTED, size=8.3)
    run = footer_p.add_run("\t")
    set_run_font(run, size=8.3, color=MUTED)
    page_label = "Page " if data["location"] == "Italy" else "Pagina "
    run = footer_p.add_run(page_label)
    set_run_font(run, size=8.3, color=MUTED)
    add_page_field(footer_p)


def add_kicker(document, text):
    paragraph = document.add_paragraph()
    paragraph_base(paragraph, before=0, after=2, line=1.0)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=8.4, color=MAROON, bold=True)
    return paragraph


def add_contact_line(document, data):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_base(paragraph, after=9, line=1.0)
    run = paragraph.add_run(f"{data['location']}  |  ")
    set_run_font(run, size=9.0, color=MUTED)
    add_hyperlink(paragraph, "lorenzo03.marinelli@gmail.com", "mailto:lorenzo03.marinelli@gmail.com", size=9.0)
    run = paragraph.add_run("  |  ")
    set_run_font(run, size=9.0, color=MUTED)
    add_hyperlink(paragraph, "Portfolio", "https://www.lmarinelli.eu", size=9.0)
    run = paragraph.add_run("  |  ")
    set_run_font(run, size=9.0, color=MUTED)
    add_hyperlink(paragraph, "GitHub", "https://github.com/Lorenzo-3", size=9.0)
    run = paragraph.add_run("  |  ")
    set_run_font(run, size=9.0, color=MUTED)
    add_hyperlink(
        paragraph,
        "LinkedIn",
        "https://www.linkedin.com/in/lorenzo-marinelli-6821403a9/",
        size=9.0,
    )


def add_profile_callout(document, data):
    add_kicker(document, data["profile_heading"])
    paragraph = document.add_paragraph()
    paragraph_base(paragraph, before=0, after=8, line=1.16, keep=True)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    set_cell_shading(paragraph, PALE_GOLD)
    set_left_border(paragraph, "C10030")
    run = paragraph.add_run(data["profile"])
    set_run_font(run, size=9.8, color=INK)


def add_section_heading(document, text):
    return document.add_paragraph(text, style="Heading 1")


def add_education(document, data):
    add_section_heading(document, data["education"])
    for title, dates, institution, result in data["education_items"]:
        paragraph = document.add_paragraph(style="CV Entry")
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        run = paragraph.add_run(title)
        set_run_font(run, size=9.7, color=TEAL, bold=True)
        run = paragraph.add_run(f"\t{dates}")
        set_run_font(run, size=8.9, color=MUTED, bold=True)
        meta = document.add_paragraph(style="CV Meta")
        run = meta.add_run(institution)
        set_run_font(run, size=8.7, color=MUTED, italic=True)
        if result:
            run = meta.add_run(f"  |  {result}")
            set_run_font(run, size=8.7, color=MAROON, bold=True)


def add_labeled_lines(document, heading, items):
    add_section_heading(document, heading)
    for label, text in items:
        paragraph = document.add_paragraph()
        paragraph_base(paragraph, after=2.7, line=1.08, keep=True)
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, size=9.3, color=TEAL, bold=True)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.3, color=INK)


def add_bullets(document, items, compact=False):
    for text in items:
        paragraph = document.add_paragraph(style="List Bullet")
        if compact:
            paragraph.paragraph_format.space_after = Pt(1.8)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.15 if compact else 9.3, color=INK)


def add_project(document, title, url, meta, bullets):
    paragraph = document.add_paragraph(style="Heading 2")
    add_hyperlink(paragraph, title, url, color=TEAL, bold=True, size=10.5)
    meta_paragraph = document.add_paragraph(meta, style="CV Meta")
    meta_paragraph.paragraph_format.keep_with_next = True
    add_bullets(document, bullets, compact=True)


def add_metric_callout(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_base(paragraph, before=2, after=5, line=1.0, keep=True)
    set_cell_shading(paragraph, PALE_TEAL)
    set_left_border(paragraph, "005A56", size="14", space="5")
    run = paragraph.add_run(text)
    set_run_font(run, size=8.7, color=TEAL, bold=True)


def add_additional_projects(document, data):
    add_section_heading(document, data["additional_projects"])
    for title, text in data["additional_project_items"]:
        paragraph = document.add_paragraph()
        paragraph_base(paragraph, after=3, line=1.1, keep=True)
        run = paragraph.add_run(f"{title}: ")
        set_run_font(run, size=9.25, color=TEAL, bold=True)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.25, color=INK)


def add_portfolio(document, data):
    add_section_heading(document, data["portfolio"])
    paragraph = document.add_paragraph()
    paragraph_base(paragraph, after=4, line=1.12, keep=True)
    run = paragraph.add_run(data["portfolio_intro"])
    set_run_font(run, size=9.5, color=INK)
    for label, text in data["portfolio_groups"]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, size=9.2, color=TEAL, bold=True)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.2, color=INK)
    link_p = document.add_paragraph()
    paragraph_base(link_p, before=1, after=5, line=1.0)
    add_hyperlink(link_p, data["playground_url"].removeprefix("https://"), data["playground_url"], bold=True, size=9.2)


def add_links(document, data):
    add_section_heading(document, data["links"])
    paragraph = document.add_paragraph()
    paragraph_base(paragraph, after=4, line=1.1, keep=True)
    run = paragraph.add_run(data["links_intro"])
    set_run_font(run, size=9.4, color=INK)
    for label, url in data["link_items"]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        add_hyperlink(paragraph, label, url, color=TEAL, bold=True, size=9.2)


def add_page_break(document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def set_document_properties(document, language):
    props = document.core_properties
    props.author = "Lorenzo Marinelli"
    props.title = f"Lorenzo Marinelli - Curriculum Vitae ({language.upper()})"
    props.subject = "Academic and project curriculum vitae"
    props.keywords = "Computer Science, Machine Learning, Bioinformatics, Computer Vision"
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def build_cv(language):
    data = CONTENT[language]
    document = Document()
    configure_styles(document)
    configure_page(document, data)
    set_document_properties(document, language)

    add_kicker(document, f"{data['cv_label']} | {data['date']}")
    document.add_paragraph("Lorenzo Marinelli", style="Title")
    document.add_paragraph(data["subtitle"], style="Subtitle")
    add_contact_line(document, data)
    add_profile_callout(document, data)
    add_education(document, data)
    add_labeled_lines(document, data["technical"], data["skill_items"])
    add_section_heading(document, data["focus"])
    add_bullets(document, data["focus_items"], compact=True)

    add_page_break(document)
    add_section_heading(document, data["projects"])
    add_project(document, data["thesis_title"], data["thesis_url"], data["thesis_meta"], data["thesis_bullets"])
    add_metric_callout(document, data["thesis_metrics"])
    add_project(document, data["rl_title"], data["rl_url"], data["rl_meta"], data["rl_bullets"])
    add_project(document, data["uav_title"], data["uav_url"], data["uav_meta"], data["uav_bullets"])
    add_project(document, data["lorenz_title"], data["lorenz_url"], data["lorenz_meta"], data["lorenz_bullets"])
    add_additional_projects(document, data)

    add_page_break(document)
    add_portfolio(document, data)
    add_section_heading(document, data["practice"])
    add_bullets(document, data["practice_items"])
    add_section_heading(document, data["interests"])
    paragraph = document.add_paragraph()
    paragraph_base(paragraph, after=4, line=1.12, keep=True)
    run = paragraph.add_run(data["interests_text"])
    set_run_font(run, size=9.5, color=INK)
    add_links(document, data)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / data["filename"]
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    for lang in ("en", "it"):
        print(build_cv(lang))
